#!/usr/bin/env python3
"""Extract all speeches by 砚仁 from meeting records (both .md and .docx)."""

import os
import re
import sys
from docx import Document

MEETING_DIR = r"c:\Users\苏砚仁\thinknote\南塘 DAO\会议记录"

def extract_from_md(filepath):
    """Extract 砚仁 speeches from a markdown meeting record."""
    speeches = []
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        print(f"  ERROR reading {filepath}: {e}", file=sys.stderr)
        return speeches

    # Find the meeting title
    title_match = re.search(r'^#\s+(.+?)$', content, re.MULTILINE)
    title = title_match.group(1) if title_match else os.path.basename(filepath)

    # Find all speech entries by 砚仁
    # Pattern: number. **砚仁** [time range] (type)
    # Followed by content until next numbered entry or end
    lines = content.split('\n')
    in_yanren = False
    current_speech = None
    speech_lines = []

    for i, line in enumerate(lines):
        # Check if this line starts a new speech
        match = re.match(r'^(\d+)\.\s+\*\*(.+?)\*\*\s+\[(.+?)\]\s*[（(](.+?)[）)]\s*$', line)
        if match:
            # Save previous speech if it was 砚仁's
            if in_yanren and current_speech:
                current_speech['text'] = '\n'.join(speech_lines).strip()
                speeches.append(current_speech)

            speaker = match.group(2).strip()
            time_range = match.group(3).strip()
            speech_type = match.group(4).strip()

            if speaker == '砚仁':
                in_yanren = True
                current_speech = {
                    'file': os.path.basename(filepath),
                    'title': title,
                    'time': time_range,
                    'type': speech_type,
                }
                speech_lines = []
            else:
                in_yanren = False
                current_speech = None
                speech_lines = []
        elif in_yanren and line.startswith('    '):
            # Continuation of 砚仁's speech (indented)
            stripped = line.strip()
            # Skip "发言整理：" and "发言要点：" sections
            if stripped.startswith('**发言整理：**') or stripped.startswith('**发言要点：**'):
                continue
            speech_lines.append(stripped)

    # Don't forget the last speech
    if in_yanren and current_speech:
        current_speech['text'] = '\n'.join(speech_lines).strip()
        speeches.append(current_speech)

    return speeches


def extract_from_docx(filepath):
    """Extract 砚仁 speeches from a docx meeting record.

    .docx format differs from .md:
    - Speaker line: 砚仁[下午8:05]自由发言  (no number, no bold markers)
    - Content on following lines
    - Then 发言要点： section
    """
    speeches = []
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"  ERROR reading {filepath}: {e}", file=sys.stderr)
        return speeches

    # Get meeting title from first paragraph
    title = doc.paragraphs[0].text if doc.paragraphs else os.path.basename(filepath)
    title = title.lstrip('#').strip()

    # Collect all paragraphs
    paragraphs = [p.text for p in doc.paragraphs]

    in_yanren = False
    current_speech = None
    speech_lines = []

    # Pattern for .docx speaker line: 砚仁[下午8:05]自由发言
    # (no number prefix, no bold markers)
    docx_speaker_re = re.compile(r'^(.+?)\[([^\]]+)\](.+?)$')

    for i, para_text in enumerate(paragraphs):
        text = para_text.strip()
        if not text:
            if in_yanren:
                speech_lines.append('')
            continue

        # Try to match speaker line: 砚仁[time]type
        match = docx_speaker_re.match(text)
        if match:
            speaker = match.group(1).strip()
            time_range = match.group(2).strip()
            speech_type = match.group(3).strip()

            # Check if this is a speech entry (speaker line, not a content line)
            # Speaker lines have specific format: short name + [time] + short type label
            # Content lines also start with names but are longer
            is_speaker_line = (
                len(speaker) <= 20 and
                len(speech_type) <= 15 and
                '：' not in speaker and
                not text.startswith('发言') and
                not text.startswith('**')
            )

            if is_speaker_line:
                # Save previous speech
                if in_yanren and current_speech:
                    current_speech['text'] = '\n'.join(speech_lines).strip()
                    # Only keep if there's actual content
                    if current_speech['text']:
                        speeches.append(current_speech)

                if speaker == '砚仁':
                    in_yanren = True
                    current_speech = {
                        'file': os.path.basename(filepath),
                        'title': title,
                        'time': time_range,
                        'type': speech_type,
                    }
                    speech_lines = []
                else:
                    in_yanren = False
                    current_speech = None
                    speech_lines = []
            elif in_yanren:
                # This matched the regex but is not a speaker line (e.g., longer content)
                # Skip 发言要点 lines
                if not text.startswith('发言要点：') and not text.startswith('发言整理：'):
                    speech_lines.append(text)
        elif in_yanren:
            # Continuation of 砚仁's speech
            # Skip 发言要点： and 发言整理： sections
            if text.startswith('发言要点：') or text.startswith('发言整理：'):
                # Also skip the numbered points that follow 发言要点
                # Those are AI-generated summaries, not original speech
                in_yanren = False
                if current_speech:
                    current_speech['text'] = '\n'.join(speech_lines).strip()
                    if current_speech['text']:
                        speeches.append(current_speech)
                    current_speech = None
                speech_lines = []
                continue
            speech_lines.append(text)

    # Don't forget the last speech
    if in_yanren and current_speech:
        current_speech['text'] = '\n'.join(speech_lines).strip()
        if current_speech['text']:
            speeches.append(current_speech)

    return speeches


def main():
    all_speeches = []

    # Process all files
    files = sorted(os.listdir(MEETING_DIR))
    md_count = 0
    docx_count = 0

    for filename in files:
        filepath = os.path.join(MEETING_DIR, filename)
        if not os.path.isfile(filepath):
            continue

        if filename.endswith('.md'):
            md_count += 1
            speeches = extract_from_md(filepath)
            if speeches:
                print(f"[MD] {filename}: {len(speeches)} speeches")
                all_speeches.extend(speeches)
        elif filename.endswith('.docx'):
            docx_count += 1
            speeches = extract_from_docx(filepath)
            if speeches:
                print(f"[DOCX] {filename}: {len(speeches)} speeches")
                all_speeches.extend(speeches)

    print(f"\nProcessed {md_count} .md files, {docx_count} .docx files")
    print(f"Total 砚仁 speeches found: {len(all_speeches)}")

    # Sort by filename (which includes date)
    all_speeches.sort(key=lambda s: s['file'])

    # Output as markdown
    output_path = os.path.join(os.path.dirname(MEETING_DIR), '砚仁发言汇总_完整版.md')
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# 苏砚仁发言资料库\n\n")
        f.write(f"> 统计范围：{len(files)} 个会议记录文件（{md_count} .md + {docx_count} .docx）\n")
        f.write(f"> 发言总数：{len(all_speeches)} 条\n")
        f.write(f"> 生成时间：2026年6月21日\n\n")
        f.write("---\n\n")

        # Group by meeting
        current_file = None
        speech_index = 0
        for s in all_speeches:
            if s['file'] != current_file:
                current_file = s['file']
                # Extract date from filename
                date_match = re.search(r'(\d{8})', current_file)
                date_str = date_match.group(1) if date_match else ''
                if date_str:
                    date_str = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"

                f.write(f"## {s['title']}\n\n")
                f.write(f"- **文件**：{current_file}\n")
                if date_str:
                    f.write(f"- **日期**：{date_str}\n")
                f.write("\n")

            speech_index += 1
            f.write(f"### 发言 {speech_index}\n\n")
            f.write(f"- **时间**：{s['time']}\n")
            f.write(f"- **类型**：{s['type']}\n\n")
            f.write(f"{s['text']}\n\n")
            f.write("---\n\n")

    print(f"\nOutput written to: {output_path}")

    # Also output a compact JSON-like summary for analysis
    summary_path = os.path.join(os.path.dirname(MEETING_DIR), '砚仁发言_纯文本.txt')
    with open(summary_path, 'w', encoding='utf-8') as f:
        for i, s in enumerate(all_speeches, 1):
            f.write(f"=== 发言 {i} | {s['file']} | {s['time']} | {s['type']} ===\n")
            f.write(s['text'] + '\n\n')

    print(f"Plain text output written to: {summary_path}")


if __name__ == '__main__':
    main()

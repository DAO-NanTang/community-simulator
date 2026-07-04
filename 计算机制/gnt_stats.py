"""
gNT 统计脚本

规则：
- 出席：参会者表中"亲自出席"即为出席，10 NT/次
- 有效发言：单次会议发言 ≥ 3 次 = 1 次有效发言，+10 NT/次
- 渠道分类：
  - 100% 渠道（新手计划、参会记录）：NT 全部可兑换成 GNT
  - 50% 渠道（组长薪资/提案/宣传等）：NT 只有一半可兑换
- 每季度汇总一次
- 铸造流程：
  1. 新 GNT = 特殊渠道 NT + 一般渠道 NT × 50%
  2. 可用 GNT = 新 GNT + 上季度保留
  3. 累积阈值：可用 GNT < 500 时不燃烧、无票权，全额保留至下季度
     可用 GNT ≥ 500 时：
     - 燃烧（票权）= 可用 GNT × 50%，票权 = √(燃烧量)
     - 下季度保留 = 可用 GNT × 50%（累积到下个季度）
"""

import re
import csv
import json
from pathlib import Path
from collections import defaultdict
from datetime import datetime

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

# ──────────────────────────────────────────────
# 配置
# ──────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
BASE_DIR = SCRIPT_DIR.parent

RECORD_DIR = BASE_DIR.parent / "会议记录"
SALARY_DIR = BASE_DIR / "工资条"

SYSTEM_ROLES = {"智能主持人", "系统", "智能秘书"}

# 只统计南塘DAO组织的会议（文件名前缀）
ORGANIZATION_PREFIX = "南塘DAO"

# ── 名字映射表（任意别名 → 标准名） ──
# 标准名（14人）：甘雨、亭真、杨云标、淑惠、张超、慢慢、刘宇、小洪、朝林、建乔、虹阳、高琳、杨振、必兵
NAME_MAPPING = {
    # --- 标准名自身（自映射） ---
    "甘雨": "甘雨",
    "亭真": "亭真",
    "杨云标": "杨云标",
    "淑惠": "淑惠",
    "张超": "张超",
    "慢慢": "慢慢",
    "刘宇": "刘宇",
    "小洪": "小洪",
    "朝林": "朝林",
    "健乔": "健乔",
    "佳彬": "佳彬",
    "虹阳": "虹阳",
    "高琳": "高琳",
    "杨振": "杨振",
    "必兵": "必兵",
    "CC": "CC",
    "KIKO": "KIKO",
    "小白": "小白",
    "跳跳": "跳跳",
    "道邓": "道邓",
    "狮蛮": "狮蛮",
    "若曦": "若曦",
    "菱形": "菱形",
    # --- 两边共有（无需映射，自映射） ---
    "世佳": "世佳",
    "余星": "余星",
    "周周": "周周",
    "家乐": "家乐",
    "淅淅": "淅淅",
    "砚仁": "砚仁",
    "菡白": "菡白",
    "言礼": "言礼",
    "柯宇": "柯宇",
    # --- 会议记录中的别名 ---
    "Sion": "甘雨",
    "YTZ": "亭真",
    "南塘汉子": "杨云标",
    "标哥": "杨云标",
    "和光同尘": "淑惠",
    "淑慧": "淑惠",
    "奇迹行者": "张超",
    "樊嘉": "慢慢",
    "木木曦": "刘宇",
    "洪水": "小洪",
    "蚊子": "朝林",
    "谠泰": "健乔",
    "建桥": "健乔",
    "鄉栝": "健乔",
    "野生蓝莓": "虹阳",
    "青峰明月": "高琳",
    "麦田": "杨振",
    "振哥": "杨振",
    "xboring": "必兵",
    "xiaobai": "小白",
    "焕炘": "道邓",
    # --- 工资表中的别名 ---
    "cc": "CC",
    "kiko": "KIKO",
    "跳": "跳跳",
}

# 排除名单（非南塘DAO成员 + 不参与 gNT 票权计算）
EXCLUDED_NAMES = {
    # 刘兵、合作社（不参与票权计算）
    "刘兵", "合作社",
    # 非南塘DAO成员
    "Daniel", "FC", "Hello 章先森", "Stella", "aa", "云展",
    "南塘之子", "文峰", "大杜", "富章", "方熊", "漫山漫野",
    "紫螟", "言礼", "严丽", "阿山", "麦客",
    "泡哥", "黄敏欢",
}

# 需要排除的无效名（如表格表头误入）
INVALID_NAMES = {"姓名"}

# 计算起始季度（此季度之前不计入 GNT 核算）
START_SEASON = "S5"


# ── 全局警告列表 ──
warnings_list: list[dict] = []  # {"type": "skipped_meeting" | "unknown_name" | "new_person", "detail": str}


def is_dao_meeting(filepath: Path) -> bool:
    """判断会议文件是否属于南塘DAO组织"""
    # 从文件名中提取组织名（格式：组织名-日期-...）
    parts = filepath.name.split("-")
    if parts:
        org_name = parts[0].strip()
        return org_name == ORGANIZATION_PREFIX
    return False


def normalize_name(name: str) -> str:
    """统一名字别名，返回标准化后的名字"""
    name = name.strip()
    if name in INVALID_NAMES:
        return ""
    return NAME_MAPPING.get(name, name)


def check_unknown_names(meetings: list, salary_data: dict) -> list:
    """检查未知名字（不在 NAME_MAPPING 中的名字），返回警告列表"""
    known_names = set(NAME_MAPPING.keys()) | set(NAME_MAPPING.values())
    unknown = set()
    for meeting in meetings:
        for name in meeting["attendees"]:
            if name not in known_names and name not in SYSTEM_ROLES and name not in EXCLUDED_NAMES:
                unknown.add(name)
        for name in meeting.get("speakers", []):
            if name not in known_names and name not in SYSTEM_ROLES and name not in EXCLUDED_NAMES:
                unknown.add(name)
    # 检查工资表中的名字
    for q, members in salary_data.items():
        for name in members:
            if name not in known_names and name not in EXCLUDED_NAMES:
                unknown.add(name)
    return sorted(unknown)


NT_ATTEND = 10
NT_EFFECTIVE_SPEECH = 10
SPEECH_THRESHOLD = 3
GNT_BURN_THRESHOLD = 500  # 累积可用 GNT 低于此值不燃烧，全额保留至下季度


# ──────────────────────────────────────────────
# 解析工资 CSV → {季度: {姓名: NT总额}}
# ──────────────────────────────────────────────
def parse_salary_csv() -> dict:
    """读取工资条 CSV，按标准季度汇总每人 NT"""
    quarterly_general = defaultdict(lambda: defaultdict(float))

    salary_files = list(SALARY_DIR.glob("*.csv"))
    if not salary_files:
        print("未找到工资条 CSV 文件。")
        return dict(quarterly_general)

    for csv_file in sorted(salary_files):
        with open(csv_file, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                nt_str = row.get("NT", "").strip()
                if not nt_str or nt_str == "0":
                    continue

                nt = float(nt_str)
                creditor = row.get("债权人", "").strip()
                quarter_raw = row.get("季度", "").strip()

                # 提取人名并统一别名
                raw_name = creditor.split("(")[0].strip() if "(" in creditor else creditor
                name = normalize_name(raw_name)
                if not name:
                    continue

                # 从季度列提取 S 编号
                season = get_season_from_label(quarter_raw)
                if not season:
                    continue

                quarterly_general[season][name] += nt

    return dict(quarterly_general)


# ──────────────────────────────────────────────
# 解析单份发言记录
# ──────────────────────────────────────────────
def parse_meeting_record(filepath: str) -> dict:
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    content_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", content)

    meta = {}
    for key, pattern in [("meeting_id", "会议编号"), ("start_time", "开始时间"), ("end_time", "结束时间")]:
        for sep in ["：", ":", ":"]:
            idx = content_clean.find(f"{pattern}{sep}")
            if idx >= 0:
                val = content_clean[idx + len(pattern) + 1:].split("\n")[0].strip().lstrip("- ").strip()
                if val:
                    meta[key] = val
                break

    date_match = re.search(r"(\d{8})", Path(filepath).stem)
    if date_match:
        meta["date"] = date_match.group(1)

    # 参会者
    attendees = {}
    if "## 参会者" in content_clean:
        table_section = content.split("## 参会者")[1]
        if "## " in table_section:
            table_section = table_section.split("## ")[0]
        for line in table_section.split("\n"):
            if "|" not in line or "---" in line:
                continue
            cols = [c.strip() for c in line.split("|") if c.strip()]
            if len(cols) >= 3:
                normalized = normalize_name(cols[0])
                if normalized:
                    attendees[normalized] = {"status": cols[2]}

    # 发言
    SPEECH_RE = re.compile(r"\d+\.\s+\*\*(.+?)\*\*\s+\[([^\]]+)\](?:\s*\(([^)]*)\))?")
    speeches = defaultdict(list)
    for line in content.split("\n"):
        m = SPEECH_RE.match(line.strip())
        if m:
            speaker = normalize_name(m.group(1).strip())
            if speaker and speaker not in SYSTEM_ROLES:
                speeches[speaker].append(m.group(2).strip())

    result = {"meta": meta, "attendees": {}}
    for name, info in attendees.items():
        if info["status"] == "亲自出席":
            speech_count = len(speeches.get(name, []))
            is_effective = speech_count >= SPEECH_THRESHOLD
            result["attendees"][name] = {
                "attended": True,
                "speech_count": speech_count,
                "effective_speech": 1 if is_effective else 0,
            }
    return result


def parse_docx_record(filepath: str) -> dict:
    """解析 .docx 格式的发言记录"""
    if not HAS_DOCX:
        raise ImportError("python-docx 未安装")

    doc = Document(filepath)
    path = Path(filepath)

    meta = {}
    date_match = re.search(r"(\d{8})", path.stem)
    if date_match:
        meta["date"] = date_match.group(1)

    # 从段落中提取元数据
    time_re = re.compile(r"(\d{4}年\d{1,2}月\d{1,2}日[上下]午\d{1,2}:\d{2})")
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("会议时间："):
            times = time_re.findall(t)
            if len(times) >= 1:
                meta["start_time"] = times[0]
            if len(times) >= 2:
                meta["end_time"] = times[1]

    attendees = {}
    speeches = defaultdict(list)

    SPEECH_RE = re.compile(r"^([一-龥a-zA-Z\s]+?)\[([^\]]+)\](.+)$")
    ATTEND_RE = re.compile(r"([一-龥a-zA-Z\s]+?)[（(]([^）)]*)[）)]")

    for p in doc.paragraphs:
        text = p.text.strip()

        # 提取出席情况
        if text.startswith("出席情况："):
            content_after = text[len("出席情况："):]
            # 按 `、` 分割每个人
            for person_str in content_after.split("、"):
                person_str = person_str.strip()
                m = ATTEND_RE.match(person_str)
                if m:
                    name = m.group(1).strip()
                    info = m.group(2)
                    status = ""
                    if "亲自出席" in info:
                        status = "亲自出席"
                    elif "缺席" in info:
                        status = "缺席"
                    elif "委托出席" in info:
                        status = "委托出席"
                    elif "旁听" in info:
                        status = "旁听"
                    if status:
                        normalized = normalize_name(name)
                        if normalized:
                            attendees[normalized] = {"status": status}

        # 提取发言记录
        m = SPEECH_RE.match(text)
        if m:
            speaker = normalize_name(m.group(1).strip())
            if speaker and speaker not in SYSTEM_ROLES:
                speeches[speaker].append(m.group(2).strip())

    result = {"meta": meta, "attendees": {}}
    for name, info in attendees.items():
        if info["status"] == "亲自出席":
            speech_count = len(speeches.get(name, []))
            is_effective = speech_count >= SPEECH_THRESHOLD
            result["attendees"][name] = {
                "attended": True,
                "speech_count": speech_count,
                "effective_speech": 1 if is_effective else 0,
            }
    return result


def parse_all_records() -> tuple[list, list[str], list[str]]:
    """解析所有南塘DAO组织的发言记录（支持 .md 和 .docx）
    Returns: (meeting_list, skipped_files, excluded_names_in_meetings)
    """
    results = []
    skipped_files = []
    excluded_in_meetings = set()

    # 先收集所有符合条件的文件，按 stem 去重
    file_candidates = {}  # stem -> preferred filepath (.md优先于.docx)
    for f in sorted(RECORD_DIR.glob("*")):
        if f.name.startswith("~") or f.is_dir():
            continue

        # 只统计南塘DAO组织的会议
        if not is_dao_meeting(f):
            skipped_files.append(f.name)
            continue

        # 跳过非发言记录文件（如会议纪要）
        if "会议纪要" in f.name:
            skipped_files.append(f.name)
            continue

        suffix = f.suffix.lower()
        if suffix not in (".md", ".docx"):
            skipped_files.append(f.name)
            continue

        stem = f.stem
        if stem not in file_candidates:
            file_candidates[stem] = f
        elif suffix == ".md":
            # .md 优先于 .docx
            old = file_candidates[stem]
            if old.suffix.lower() == ".docx":
                file_candidates[stem] = f

    for stem, f in file_candidates.items():
        suffix = f.suffix.lower()
        if suffix == ".md":
            parser = parse_meeting_record
        else:
            parser = parse_docx_record
        try:
            meeting = parser(str(f))
            # 从文件名提取日期和会议编号
            _enrich_meta_from_filename(meeting, stem)
            results.append(meeting)
        except Exception as e:
            print(f"解析失败: {f.name} — {e}")

    return results, skipped_files, sorted(excluded_in_meetings)


def _enrich_meta_from_filename(meeting: dict, stem: str):
    """从文件名补充元数据（日期、会议编号）"""
    date_match = re.search(r"(\d{8})", stem)
    if date_match and "date" not in meeting["meta"]:
        meeting["meta"]["date"] = date_match.group(1)

    # 从文件名提取会议编号（最后一个数字段）
    num_match = re.search(r"-(\d+)$", stem)
    if num_match and "meeting_id" not in meeting["meta"]:
        meeting["meta"]["meeting_id"] = num_match.group(1)


# ──────────────────────────────────────────────
# 季度分组
# ──────────────────────────────────────────────
def get_quarter(date_str: str) -> str:
    try:
        dt = datetime.strptime(date_str, "%Y%m%d")
        month = dt.month
        year = dt.year
        # 以 2025年 = 基准年, Q4 2025 = S4, 每年 4 个季度
        base = (year - 2025) * 4
        if month >= 10:
            return f"S{base + 4}"
        elif month >= 7:
            return f"S{base + 3}"
        elif month >= 4:
            return f"S{base + 2}"
        else:
            return f"S{base + 1}"
    except ValueError:
        return "未知季度"


def get_season_from_label(label: str) -> str:
    """从工资 CSV 季度列提取 S 编号，如 'S5季度1月份' → 'S5'"""
    m = re.search(r"(S\d+)", label)
    return m.group(1) if m else ""


# ──────────────────────────────────────────────
# 按季度汇总（合并会议 + 工资）
# ──────────────────────────────────────────────
def aggregate_data(meetings: list, salary_data: dict) -> dict:
    # 会议数据
    quarterly = defaultdict(lambda: defaultdict(lambda: {
        "attend_count": 0,
        "effective_speech_count": 0,
        "meetings": [],
    }))

    for meeting in meetings:
        meta = meeting["meta"]
        date = meta.get("date", "unknown")
        q = get_quarter(date)

        for name, data in meeting["attendees"].items():
            entry = quarterly[q][name]
            entry["attend_count"] += 1
            entry["effective_speech_count"] += data["effective_speech"]
            meeting_label = f"{meta.get('start_time', date)}"
            if data["effective_speech"]:
                meeting_label += " ★"
            entry["meetings"].append(meeting_label)

    # 合并工资数据（一般渠道 NT）
    for q, members in quarterly.items():
        salary_q = salary_data.get(q, {})
        # 加入工资数据中但无会议记录的人
        for name in salary_q:
            if name not in members:
                members[name] = {
                    "attend_count": 0,
                    "effective_speech_count": 0,
                    "meetings": [],
                }
        for name in members:
            attend_nt = members[name]["attend_count"] * NT_ATTEND
            speech_nt = members[name]["effective_speech_count"] * NT_EFFECTIVE_SPEECH
            members[name]["special_nt"] = attend_nt + speech_nt
            members[name]["general_nt"] = salary_q.get(name, 0)
            members[name]["total_nt"] = members[name]["special_nt"] + members[name]["general_nt"]

    # 确保工资数据中的季度都有条目（即使没有会议记录）
    for q, salary_members in salary_data.items():
        if q not in quarterly:
            for name, general_nt in salary_members.items():
                quarterly[q][name] = {
                    "attend_count": 0,
                    "effective_speech_count": 0,
                    "special_nt": 0,
                    "general_nt": general_nt,
                    "total_nt": general_nt,
                    "meetings": [],
                }
        else:
            # 季度已存在（有会议记录），补入工资数据
            for name, general_nt in salary_members.items():
                if name not in quarterly[q]:
                    quarterly[q][name] = {
                        "attend_count": 0,
                        "effective_speech_count": 0,
                        "special_nt": 0,
                        "general_nt": general_nt,
                        "total_nt": general_nt,
                        "meetings": [],
                    }

    # 移除排除名单
    for q in quarterly:
        for excl in EXCLUDED_NAMES:
            quarterly[q].pop(excl, None)

    return dict(quarterly)


# ──────────────────────────────────────────────
# gNT 核算（跨季度累积）
# ──────────────────────────────────────────────
def calculate_gnt(quarterly: dict) -> dict:
    results = {}
    start_num = int(START_SEASON[1:])  # S5 → 5

    def season_num(s: str) -> int:
        m = re.match(r"S(\d+)", s)
        return int(m.group(1)) if m else 0

    sorted_quarters = sorted(
        (q for q in quarterly.keys() if season_num(q) >= start_num)
    )
    retained_balance = defaultdict(float)

    for quarter in sorted_quarters:
        members = quarterly[quarter]
        results[quarter] = {}

        for name in sorted(members.keys()):
            data = members[name]
            special_nt = data["special_nt"]
            general_nt = data["general_nt"]

            new_gnt = special_nt + general_nt * 0.5
            prev_retained = retained_balance[name]
            available_gnt = new_gnt + prev_retained

            if available_gnt < GNT_BURN_THRESHOLD:
                # 累积不足阈值：不燃烧、无票权，本季度新增 GNT 作废
                # 仅有上季保留继续滚存至下季度
                burned = 0
                retained = prev_retained  # 本季度新增不累积
                voting_power = 0
            else:
                burned = available_gnt * 0.5
                retained = available_gnt * 0.5
                voting_power = burned ** 0.5 if burned > 0 else 0

            retained_balance[name] = retained

            results[quarter][name] = {
                "special_nt": special_nt,
                "general_nt": general_nt,
                "new_gnt": new_gnt,
                "prev_retained": prev_retained,
                "available_gnt": available_gnt,
                "burned": burned,
                "retained": retained,
                "voting_power": round(voting_power, 2),
            }

    return results


# ──────────────────────────────────────────────
# 输出
# ──────────────────────────────────────────────
def print_report(meetings: list, quarterly: dict, gnt_results: dict):
    print("=" * 70)
    print("  gNT 参会统计报告")
    print("=" * 70)

    # 单次会议详情
    print("\n—— 单次会议详情 ——\n")
    for meeting in meetings:
        meta = meeting["meta"]
        start = meta.get("start_time", "unknown")
        m_id = meta.get("meeting_id", "")
        print(f"会议: {start}  (编号: {m_id})")
        print(f"  {'姓名':<12} {'出席':>4} {'发言次数':>6} {'有效发言':>6} {'NT':>6}")
        print(f"  {'─'*50}")
        for name, data in sorted(meeting["attendees"].items()):
            eff = data["effective_speech"]
            effective = "是" if eff else "否"
            nt = NT_ATTEND + (NT_EFFECTIVE_SPEECH if eff else 0)
            print(f"  {name:<12} {data['attended']:>4} {data['speech_count']:>6} {effective:>6} {nt:>6}")
        print()

    # 季度汇总
    start_num = int(START_SEASON[1:])
    def season_num(s: str) -> int:
        m = re.match(r"S(\d+)", s)
        return int(m.group(1)) if m else 0

    print("\n—— 季度汇总 ——\n")
    for quarter in sorted(quarterly.keys()):
        if season_num(quarter) < start_num:
            continue
        members = quarterly[quarter]
        gnt = gnt_results.get(quarter, {})
        print(f"  季度: {quarter}")
        print()
        print(f"  {'姓名':<10} {'参会':>4} {'有效发言':>6} "
              f"{'特殊NT':>8} {'一般NT':>8} "
              f"{'可兑GNT':>8} {'上季保留':>8} {'可用GNT':>8} "
              f"{'燃烧':>8} {'保留':>8} {'票权':>6}")
        print(f"  {'─'*115}")
        for name in sorted(members.keys()):
            d = members[name]
            g = gnt.get(name, {})
            print(f"  {name:<10} {d['attend_count']:>4} {d['effective_speech_count']:>6} "
                  f"{d['special_nt']:>8.0f} {d['general_nt']:>8.0f} "
                  f"{g.get('new_gnt', 0):>8.0f} {g.get('prev_retained', 0):>8.0f} "
                  f"{g.get('available_gnt', 0):>8.0f} "
                  f"{g.get('burned', 0):>8.0f} {g.get('retained', 0):>8.0f} "
                  f"{g.get('voting_power', 0):>6}")
        print()


def save_json(data, filepath: str):
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"已保存: {filepath}")


HEADER_FONT = Font(name="微软雅黑", bold=True, size=10, color="FFFFFF")
HEADER_FILL = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
TITLE_FONT = Font(name="微软雅黑", bold=True, size=12)
NORMAL_FONT = Font(name="微软雅黑", size=10)
THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)
CENTER = Alignment(horizontal="center", vertical="center")
LEFT_ALIGN = Alignment(horizontal="left", vertical="center")


def _style_sheet(ws, headers: list, rows: list, col_widths: list | None = None):
    """Apply consistent formatting to a worksheet"""
    # Title row
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    title_cell = ws.cell(row=1, column=1, value=ws.title)
    title_cell.font = TITLE_FONT
    title_cell.alignment = CENTER

    # Header row
    for col_idx, h in enumerate(headers, start=1):
        cell = ws.cell(row=2, column=col_idx, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = CENTER
        cell.border = THIN_BORDER

    # Data rows
    for row_idx, row_data in enumerate(rows, start=3):
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.font = NORMAL_FONT
            cell.border = THIN_BORDER
            if col_idx == 1:
                cell.alignment = LEFT_ALIGN
            else:
                cell.alignment = CENTER

    # Column widths
    if col_widths:
        for col_idx, w in enumerate(col_widths, start=1):
            ws.column_dimensions[chr(64 + col_idx)].width = w


def export_to_excel(
    meetings: list,
    quarterly: dict,
    gnt_results: dict,
    filepath: str,
):
    """导出 gNT 计算结果到 Excel 文件"""
    if not HAS_XLSX:
        print("⚠ openpyxl 未安装，跳过 Excel 导出。")
        return

    wb = Workbook()
    wb.remove(wb.active)  # Remove default sheet

    start_num = int(START_SEASON[1:])

    def season_num(s: str) -> int:
        m = re.match(r"S(\d+)", s)
        return int(m.group(1)) if m else 0

    # ── Sheet 1: S5季度核算 ──
    ws = wb.create_sheet("S5季度核算")
    headers = ["姓名", "参会次数", "有效发言", "特殊NT", "一般NT", "可兑GNT", "上季保留", "可用GNT", "燃烧", "保留", "票权"]
    rows = []
    members = quarterly.get("S5", {})
    gnt = gnt_results.get("S5", {})
    for name in sorted(members.keys()):
        d = members[name]
        g = gnt.get(name, {})
        rows.append([
            name,
            d["attend_count"],
            d["effective_speech_count"],
            d["special_nt"],
            d["general_nt"],
            g.get("new_gnt", 0),
            g.get("prev_retained", 0),
            g.get("available_gnt", 0),
            g.get("burned", 0),
            g.get("retained", 0),
            g.get("voting_power", 0),
        ])
    _style_sheet(ws, headers, rows, [10, 8, 8, 10, 10, 10, 10, 10, 10, 10, 8])

    # ── Sheet 2: S6季度核算 ──
    ws = wb.create_sheet("S6季度核算")
    members = quarterly.get("S6", {})
    gnt = gnt_results.get("S6", {})
    rows = []
    for name in sorted(members.keys()):
        d = members[name]
        g = gnt.get(name, {})
        rows.append([
            name,
            d["attend_count"],
            d["effective_speech_count"],
            d["special_nt"],
            d["general_nt"],
            g.get("new_gnt", 0),
            g.get("prev_retained", 0),
            g.get("available_gnt", 0),
            g.get("burned", 0),
            g.get("retained", 0),
            g.get("voting_power", 0),
        ])
    _style_sheet(ws, headers, rows, [10, 8, 8, 10, 10, 10, 10, 10, 10, 10, 8])

    # ── Sheet 3: 季度汇总（所有季度合并） ──
    ws = wb.create_sheet("季度汇总")
    summary_headers = ["姓名", "季度", "参会次数", "有效发言", "特殊NT", "一般NT", "可兑GNT", "上季保留", "可用GNT", "燃烧", "保留", "票权"]
    summary_rows = []
    for quarter in sorted(quarterly.keys()):
        if season_num(quarter) < start_num:
            continue
        members = quarterly[quarter]
        gnt = gnt_results.get(quarter, {})
        for name in sorted(members.keys()):
            d = members[name]
            g = gnt.get(name, {})
            summary_rows.append([
                name, quarter,
                d["attend_count"],
                d["effective_speech_count"],
                d["special_nt"],
                d["general_nt"],
                g.get("new_gnt", 0),
                g.get("prev_retained", 0),
                g.get("available_gnt", 0),
                g.get("burned", 0),
                g.get("retained", 0),
                g.get("voting_power", 0),
            ])
    _style_sheet(ws, summary_headers, summary_rows, [10, 8, 8, 8, 10, 10, 10, 10, 10, 10, 10, 8])

    # ── Sheet 4: 会议明细（每次会议的逐人记录） ──
    ws = wb.create_sheet("会议明细")
    mtg_headers = ["会议时间", "会议编号", "姓名", "出席", "发言次数", "有效发言", "单次会议NT"]
    mtg_rows = []
    for meeting in meetings:
        meta = meeting["meta"]
        start_time = meta.get("start_time", "")
        m_id = meta.get("meeting_id", "")
        for name, data in sorted(meeting["attendees"].items()):
            eff = data["effective_speech"]
            nt = NT_ATTEND + (NT_EFFECTIVE_SPEECH if eff else 0)
            mtg_rows.append([
                start_time,
                m_id,
                name,
                "出席",
                data["speech_count"],
                "是" if eff else "否",
                nt,
            ])
    _style_sheet(ws, mtg_headers, mtg_rows, [22, 14, 10, 8, 8, 8, 12])

    wb.save(filepath)
    print(f"已保存: {filepath}")


# ──────────────────────────────────────────────
# 主入口
# ──────────────────────────────────────────────
def main():
    meetings, skipped_files, _ = parse_all_records()
    if not meetings:
        print("未找到任何发言记录文件。")
        return

    salary_data = parse_salary_csv()

    # 警告：被跳过的非南塘DAO会议文件
    if skipped_files:
        print(f"⚠ 以下会议记录不属于南塘DAO组织，已跳过：")
        for f in skipped_files:
            print(f"  - {f}")
        print()

    # 检查未知名字（不在 NAME_MAPPING 中的名字）
    unknown_names = check_unknown_names(meetings, salary_data)
    if unknown_names:
        print("⚠ 发现以下名字未在名字映射表中，请确认：")
        for name in unknown_names:
            print(f"  - {name}")
        print()
        print("提示：这些名字尚未纳入计算，请先确认：")
        print("  1. 是否为南塘DAO成员？")
        print("  2. 如果是，其对应的标准名是什么？（若与别名相同则不需要映射）")
        print()
        return  # 等待用户确认后再继续

    print(f"共解析 {len(meetings)} 份会议记录")
    if salary_data:
        total_q = len(salary_data)
        total_persons = sum(len(v) for v in salary_data.values())
        print(f"共解析 {total_q} 个季度的工资数据（{total_persons} 人次）\n")
    else:
        print()

    quarterly = aggregate_data(meetings, salary_data)
    gnt_results = calculate_gnt(quarterly)

    print_report(meetings, quarterly, gnt_results)

    output = {
        "meeting_details": [],
        "quarterly": {},
        "gnt_results": {},
    }

    for m in meetings:
        output["meeting_details"].append({"meta": m["meta"], "attendees": m["attendees"]})

    start_num = int(START_SEASON[1:])
    def season_num(s: str) -> int:
        m2 = re.match(r"S(\d+)", s)
        return int(m2.group(1)) if m2 else 0

    for q, members in quarterly.items():
        if season_num(q) >= start_num:
            output["quarterly"][q] = dict(members)

    output["gnt_results"] = gnt_results

    json_path = SCRIPT_DIR / "gnt_stats.json"
    save_json(output, str(json_path))

    # 导出 Excel
    if HAS_XLSX:
        now = datetime.now()
        xlsx_name = f"gnt_stats_{now:%Y%m%d_%H%M}.xlsx"
        xlsx_path = BASE_DIR / "计算结果" / xlsx_name
        export_to_excel(meetings, quarterly, gnt_results, str(xlsx_path))


if __name__ == "__main__":
    main()
